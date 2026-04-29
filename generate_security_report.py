from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_security_report():
    doc = Document()
    
    # Title
    title = doc.add_heading('BunkVauth: Next-Generation SSO Security', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Security Analysis and Comparison with Traditional SSO Systems')
    
    # Section 1
    doc.add_heading('1. The Shift to Zero-Trust Authentication', level=1)
    doc.add_paragraph(
        'Traditional SSO creates a single point of failure. BunkVauth addresses this by '
        'moving from static, one-time authentication to dynamic, continuous verification.'
    )
    
    # Comparison Table
    doc.add_heading('Security Comparison Table', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Feature'
    hdr_cells[1].text = 'Traditional SSO'
    hdr_cells[2].text = 'BunkVauth (Improved)'
    
    comparisons = [
        ('Authentication Type', 'One-time (Login only)', 'Continuous (Real-time monitoring)'),
        ('Credential Type', 'Passwords / SMS OTP', 'FIDO2 Biometrics (Unphishable)'),
        ('Session Security', 'Binary (Trusted until expiry)', 'Risk-based (Zero-Trust scoring)'),
        ('Theft Protection', 'Vulnerable to Cookie Theft', 'Protected by Behavioral Signature'),
        ('User Experience', 'Interrupted by repeated prompts', 'Seamless background verification')
    ]
    
    for f, t, b in comparisons:
        row = table.add_row().cells
        row[0].text = f
        row[1].text = t
        row[2].text = b

    # Section 2
    doc.add_heading('2. Key Security Pillars', level=1)
    
    pillars = [
        ('FIDO2/WebAuthn Integration', 'Eliminates passwords entirely, using hardware-backed biometrics that are mathematically impossible to phish.'),
        ('Continuous Behavioral Biometrics', 'Captures "How you act" rather than just "What you know". This creates a second layer of security that follows the user throughout the session.'),
        ('AI Risk Scoring Engine', 'Uses Random Forest ML to calculate an anomaly score. This allows the system to detect account takeover in under 5 seconds.')
    ]
    
    for p_title, p_desc in pillars:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(p_title + ': ')
        run.bold = True
        p.add_run(p_desc)

    # Section 3
    doc.add_heading('3. Summary for Project Defense', level=1)
    summary = doc.add_paragraph(
        'BunkVauth transforms the SSO model into a proactive security framework. By '
        'implementing Hybrid Continuous Authentication, it effectively mitigates '
        'Session Hijacking and Credential Theft, making it a robust solution for '
        'enterprise-grade identity management.'
    )
    
    doc.save('BunkVauth_Security_Analysis.docx')
    print("Report generated: BunkVauth_Security_Analysis.docx")

if __name__ == "__main__":
    create_security_report()
